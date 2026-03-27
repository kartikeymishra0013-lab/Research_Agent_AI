"""
Word document (.docx) loader using python-docx.
Extracts paragraphs, tables, headers, and document properties.
"""
from __future__ import annotations

from pathlib import Path

from src.ingestion.base_loader import BaseLoader, Document
from src.utils.logger import get_logger

logger = get_logger(__name__)


class DocxLoader(BaseLoader):
    """Loads .docx files and extracts structured text content."""

    def __init__(self, include_tables: bool = True, include_headers: bool = True):
        self.include_tables = include_tables
        self.include_headers = include_headers

    def load(self, source: str) -> Document:
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn

        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {source}")

        logger.info(f"Loading DOCX: {path.name}")
        doc = DocxDocument(str(path))

        sections = []

        # Extract paragraphs (including headings)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else "Normal"
            if "Heading" in style and self.include_headers:
                sections.append(f"\n## {text}\n")
            else:
                sections.append(text)

        # Extract tables
        if self.include_tables:
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    sections.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]\n")

        full_text = "\n".join(sections)

        # Extract core properties
        props = doc.core_properties
        metadata = {
            "title": props.title or path.stem,
            "author": props.author or "",
            "subject": props.subject or "",
            "keywords": props.keywords or "",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "file_size_bytes": path.stat().st_size,
        }
        if props.created:
            metadata["created"] = props.created.isoformat()
        if props.modified:
            metadata["modified"] = props.modified.isoformat()

        logger.info(f"Extracted {len(full_text)} chars from DOCX ({len(doc.paragraphs)} paragraphs)")
        return Document(
            content=full_text,
            source=str(path),
            doc_type="docx",
            metadata=metadata,
        )
